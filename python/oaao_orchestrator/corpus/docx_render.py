"""Markdown / blocks → DOCX (CS-3-S2)."""

from __future__ import annotations

import io
import re
from typing import Any

from docx import Document
from docx.shared import Pt

from oaao_orchestrator.library.blocks import blocks_to_markdown


def markdown_to_docx_bytes(markdown: str, *, title: str = "") -> bytes:
    doc = Document()
    if title.strip():
        doc.add_heading(title.strip(), level=0)

    for raw_line in (markdown or "").splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            continue
        if re.match(r"^[-*+]\s+", line):
            doc.add_paragraph(re.sub(r"^[-*+]\s+", "", line), style="List Bullet")
            continue
        if re.match(r"^\d+\.\s+", line):
            doc.add_paragraph(re.sub(r"^\d+\.\s+", "", line), style="List Number")
            continue
        if line.startswith("```"):
            continue
        p = doc.add_paragraph(line)
        for run in p.runs:
            run.font.size = Pt(11)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def blocks_to_docx_bytes(blocks: list[dict[str, Any]], *, title: str = "") -> bytes:
    md = blocks_to_markdown(blocks, title=title)
    return markdown_to_docx_bytes(md, title=title)
